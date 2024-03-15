from __future__ import (absolute_import, division, print_function,
                        unicode_literals)

import os

import docx
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from hexs_rag.model.model.paragraph import Paragraph as ParagraphHexa


class WordReader:
    """
    A class that extrats content from a .docx file and returns a list of Paragraph objects. 
    The behaviour is straightforward: it iterates through the paragraphs and tables of the document,
    extracting the text and the style of each one. The style is determined by the style of the paragraph 
    or the predominant style of the table. The text is extracted from the paragraphs and the tables are
    converted to a string. The class also estimates the page number of each paragraph based on the total
    number of characters in the document.

    Attributes:
    path (str): The path to the .docx file.
    paragraphs (list[ParagraphHexa]): A list of Paragraph objects.

    Methods:
    iter_block_items(parent): Iterates through the paragraphs and tables of the document.
    get_paragraphs(): Extracts the paragraphs and tables from the document.
    determine_predominant_style(styles): Determines the predominant style of a list of styles.
    estimate_page_number(total_characters): Estimates the page number of a paragraph based on the total number of characters.
    extract_paragraph_info(paragraph): Extracts the text and style of a paragraph.
    table_to_paragraph(table): Converts a table to a string and determines the predominant style.
    print_paragraphs_and_tables(): Prints the paragraphs and tables of the document.
    """

    def __init__(self, path):
        self.path = path
        self.paragraphs = self.get_paragraphs()

    def iter_block_items(self, parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Unsupported parent type")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def get_paragraphs(self):
        if not os.path.exists(self.path):
            raise FileNotFoundError(f"The file {self.path} does not exist.")
        try:
            doc = docx.Document(self.path)
            paragraph_objects = []
            paragraph_id = 0
            page_id = 1  # Example page ID
            total_characters = 0
            for block in self.iter_block_items(doc):
                if isinstance(block, Paragraph):
                    paragraph_info = self.extract_paragraph_info(block)
                    if paragraph_info:  # Only append if paragraph is not empty
                        page_id = self.estimate_page_number(total_characters)
                        p_obj = ParagraphHexa(
                            text=paragraph_info["text"],
                            font_style=paragraph_info["style"],
                            id_=paragraph_id,
                            page_id=page_id,
                        )
                        # print(f"Found paragraph: {paragraph_info['style']}...")  # DEBUG
                        paragraph_objects.append(p_obj)
                        paragraph_id += 1
                        total_characters += len(paragraph_info["text"])
                elif isinstance(block, Table):
                    table_paragraph, table_style = self.table_to_paragraph(block)
                    if table_paragraph.strip():  # Check if table paragraph is not empty
                        # print(f"Found table. Predominant style: {table_style}")  # DEBUG
                        p_obj = ParagraphHexa(
                            text=table_paragraph,
                            font_style=table_style,
                            id_=paragraph_id,
                            page_id=page_id,
                        )
                        paragraph_objects.append(p_obj)
                        paragraph_id += 1
            return paragraph_objects
        except Exception as e:
            raise ValueError(f"Error reading the .docx file. Original error: {str(e)}")

    @staticmethod
    def determine_predominant_style(styles):
        # Count the occurrences of each style
        style_counts = {}
        for style in styles:
            if style in style_counts:
                style_counts[style] += 1
            else:
                style_counts[style] = 1

        # Find the style with the highest count
        predominant_style = max(style_counts, key=style_counts.get, default="None")
        if predominant_style == "Table Paragraph":
            predominant_style = "Body Text"
        return predominant_style

    @staticmethod
    def estimate_page_number(total_characters):
        avg_chars_per_page = 2000
        return total_characters // avg_chars_per_page + 1

    @staticmethod
    def extract_paragraph_info(paragraph):
        # Check if paragraph is empty
        if not paragraph.text.strip():
            return None  # Return None for empty paragraphs

        paragraph_style = paragraph.style.name if paragraph.style else "None"
        if paragraph_style == "Normal":
            paragraph_style = "Body Text"

        # Split the text into smaller paragraphs
        max_paragraph_length = 1500  # Set the maximum length of a paragraph
        paragraphs = [
            paragraph.text[i : i + max_paragraph_length]
            for i in range(0, len(paragraph.text), max_paragraph_length)
        ]

        runs = []
        for p in paragraphs:
            for run in paragraph.runs:
                run_details = {
                    "text": p,
                    "font_name": run.font.name,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                }
                runs.append(run_details)
        return {"text": paragraph.text, "style": paragraph_style, "runs": runs}

    def table_to_paragraph(self, table):
        table_text = ""
        table_styles = set()

        for row in table.rows:
            for cell in row.cells:
                cell_text = ""
                for paragraph in cell.paragraphs:
                    paragraph_style = (
                        paragraph.style.name if paragraph.style else "None"
                    )
                    table_styles.add(paragraph_style)

                    for run in paragraph.runs:
                        cell_text += run.text

                    cell_text += " "
                table_text += cell_text.strip() + " | "  # Add a separator for cells
            table_text = table_text.strip() + "\n"  # Add a newline for rows

        predominant_style = self.determine_predominant_style(table_styles)

        return table_text.strip(), predominant_style

    def print_paragraphs_and_tables(self):
        try:

            doc_items = self.get_paragraphs()
            for item in doc_items:
                if "paragraph" in item:
                    print("Paragraph:", item["paragraph"]["text"])
                elif "table" in item:
                    print("Table:")
                    for row in item["table"]:
                        for cell in row:
                            for paragraph in cell:
                                print("Cell Paragraph:", paragraph["text"])
                print("-" * 40)  # separator for clarity

        except Exception as e:
            print(f"Error: {str(e)}")
