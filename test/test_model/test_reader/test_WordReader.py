from hexs_rag.model.readers import WordReader
import pytest 
from unittest.mock import MagicMock
from docx.shared import Pt
from unittest.mock import patch


# Test Initialization
def test_init(word_reader, sample_docx_path):
    assert word_reader.path == sample_docx_path
    assert len(word_reader.paragraphs) > 0
    assert word_reader.path == sample_docx_path
    assert isinstance(word_reader.paragraphs, list)

#determine_predominant_style
def test_determine_predominant_style():
    styles = ["Body Text", "Heading", "Heading", "Body Text", "Body Text"]
    assert WordReader.determine_predominant_style(styles) == "Body Text"

def test_determine_predominant_style_with_empty_list():
    assert WordReader.determine_predominant_style([]) == "None"

def test_iter_block_items_with_unsupported_parent_type(word_reader):
    with pytest.raises(ValueError):
        list(word_reader.iter_block_items("unsupported type"))

def test_extract_paragraph_info_non_empty_paragraph(word_reader):
    paragraph = MagicMock()
    paragraph.text = "Test paragraph"
    paragraph.style.name = "Body Text"
    # Mock runs to simulate text formatting details
    run = MagicMock()
    run.text = "Test paragraph"
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = False
    run.italic = False
    run.underline = False
    paragraph.runs = [run]

    info = word_reader.extract_paragraph_info(paragraph)
    assert info is not None
    assert info['text'] == "Test paragraph"
    assert info['style'] == "Body Text"
    assert isinstance(info['runs'], list)

def test_table_to_paragraph_empty_table(word_reader):
    table = MagicMock()
    table.rows = []
    text, style = word_reader.table_to_paragraph(table)
    assert text == ""
    assert style == "None"

def test_get_paragraphs_exception_handling(sample_docx_path):
    with patch('docx.Document', side_effect=Exception("Mocked exception")):
        with pytest.raises(ValueError):
            WordReader(sample_docx_path)

def test_determine_predominant_style_with_tie():
    styles = ["Heading", "Body Text", "Heading", "Body Text"]
    # Assuming a deterministic outcome for a tie; this might need to be adjusted based on actual implementation logic.
    expected_style = "Heading" 
    assert WordReader.determine_predominant_style(styles) == expected_style

def test_determine_predominant_style_rare_styles():
    styles = ["RareStyle1", "RareStyle2", "CommonStyle", "CommonStyle"]
    assert WordReader.determine_predominant_style(styles) == "CommonStyle"

def test_extract_paragraph_info_with_complex_formatting(word_reader):
    paragraph = MagicMock()
    paragraph.text = "Complex formatting: Bold, Italic, and Underline."
    paragraph.style.name = "Complex Style"
    # Mock runs with different formatting
    run_bold = MagicMock(text="Bold", font=MagicMock(name="Arial", size=Pt(12), bold=True), italic=False, underline=False)
    run_italic = MagicMock(text="Italic", font=MagicMock(name="Arial", size=Pt(12), italic=True), bold=False, underline=False)
    run_underline = MagicMock(text="Underline", font=MagicMock(name="Arial", size=Pt(12), underline=True), bold=False, italic=False)
    paragraph.runs = [run_bold, run_italic, run_underline]

    info = word_reader.extract_paragraph_info(paragraph)
    assert info is not None
    assert len(info['runs']) == 3  # Expecting 3 runs for the mixed formatting

def test_paragraphs_with_real_file(word_reader_with_actual_file):
    # check length of paragraph list
    assert len(word_reader_with_actual_file.paragraphs) == 125
    # check text content
    assert word_reader_with_actual_file.paragraphs[0].text == 'Template for Preparation of Papers for IEEE Sponsored Conferences & Symposia'
    assert word_reader_with_actual_file.paragraphs[1].text == 'Frank Anderson, Sam B. Niles, Jr., and Theodore C. Donald, Member, IEEE'
    # check ids
    assert word_reader_with_actual_file.paragraphs[0].id_ == 210
    assert word_reader_with_actual_file.paragraphs[1].id_ == 211
    # check font style
    assert word_reader_with_actual_file.paragraphs[0].font_style == 'Title'
    assert word_reader_with_actual_file.paragraphs[1].font_style == 'Authors'
    # page id
    assert word_reader_with_actual_file.paragraphs[0].page_id == 1

@pytest.mark.skip(reason='Do later')
def test_table_to_paragraph_complex_structure(word_reader):
    table = MagicMock()
    row1 = MagicMock()
    cell1_row1 = MagicMock()
    cell2_row1 = MagicMock()
    paragraph1_cell1_row1 = MagicMock(text="Row 1, Cell 1", style=MagicMock(name="Body Text"))
    paragraph1_cell2_row1 = MagicMock(text="Row 1, Cell 2", style=MagicMock(name="Heading"))
    cell1_row1.paragraphs = [paragraph1_cell1_row1]
    cell2_row1.paragraphs = [paragraph1_cell2_row1]
    row1.cells = [cell1_row1, cell2_row1]

    # Simulate a second row with different content and style
    row2 = MagicMock()
    cell1_row2 = MagicMock()
    cell2_row2 = MagicMock()
    paragraph1_cell1_row2 = MagicMock(text="Row 2, Cell 1", style=MagicMock(name="Caption"))
    paragraph1_cell2_row2 = MagicMock(text="", style=MagicMock(name="Heading"))  # Empty cell
    cell1_row2.paragraphs = [paragraph1_cell1_row2]
    cell2_row2.paragraphs = [paragraph1_cell2_row2]
    row2.cells = [cell1_row2, cell2_row2]
    table.rows = [row1, row2]
    text, style = word_reader.table_to_paragraph(table)
    assert "Row 1, Cell 1 | Row 1, Cell 2\nRow 2, Cell 1 | " in text
    expected_style = "Body Text"  # Assuming Body Text is considered predominant for this test
    assert style == expected_style

def test_get_paragraphs_file_not_found():
    with pytest.raises(FileNotFoundError):
        WordReader("non_existent_file.docx")

def test_get_paragraphs_with_permission_error(sample_docx_path):
    with patch("builtins.open", side_effect=PermissionError("Mocked permission error")):
        with pytest.raises(Exception):
            WordReader(sample_docx_path)

def test_get_paragraphs_with_corrupted_file(sample_docx_path):
    # Assuming the existence of a mocked function that simulates opening a corrupted file
    with patch('docx.Document', side_effect=ValueError("Mocked corruption error")):
        with pytest.raises(ValueError):
            WordReader(sample_docx_path)


