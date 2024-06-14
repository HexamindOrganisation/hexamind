import pytest
from docx import Document as DocxDocument
from hexamind.model.readers.WordReader import WordReader
import os

@pytest.fixture
def create_word_doc(tmpdir):
    def _create_word_doc(content, file_name='test.docx'):
        file_path = tmpdir.join(file_name)
        doc = DocxDocument()
        for item in content:
            if item[0] == 'heading':
                doc.add_heading(item[1], level=item[2])
            elif item[0] == 'paragraph':
                doc.add_paragraph(item[1])
            elif item[0] == 'table':
                table = doc.add_table(rows=len(item[1]), cols=len(item[1][0]))
                for row_idx, row_data in enumerate(item[1]):
                    row = table.rows[row_idx]
                    for col_idx, cell_data in enumerate(row_data):
                        row.cells[col_idx].text = cell_data
        doc.save(file_path)
        return str(file_path)
    return _create_word_doc

def test_word_to_markdown_conversion(create_word_doc):
    content = [
        ('heading', 'Heading 1', 1),
        ('paragraph', 'Some paragraph.'),
        ('table', [['A', 'B'], ['1', '2']])
        ('heading', 'Heading 2', 2),
        ('paragraph', 'Another paragraph')
    ]
    file_path = create_word_doc(content)
    reader = WordReader(file_path)
    expected_markdown = '# Heading 1\n\nSome paragraph.\n\nA | B |\n|---|---|\n| 1 | 2\n\n## Heading 2\n\nAnother paragraph\n\n'
    assert reader.convert_to_markdown() == expected_markdown