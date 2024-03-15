#Description: define fixtures used within other tests

import pytest
import docx

from hxm_rag.model.model.doc import Doc
from hxm_rag.model.readers import WordReader
from hxm_rag.model.readers.HTMLreader import HtmlReader

########HTML
@pytest.fixture
def html_reader_instance(): # real html file example
    html_file_path = "data/test_data/HTML5 Test Page.html"
    return HtmlReader(html_file_path)

@pytest.fixture
def sample_html_file(tmp_path): # sample html file example
    content = """
    <!DOCTYPE html>
    <html>
    <head><title>Test Document</title></head>
    <body>
        <h1>Document Heading</h1>
        <p>This is a sample paragraph.</p>
        <ul>
            <li>List item 1</li>
            <li>List item 2</li>
        </ul>
        <table>
            <tr><th>Header 1</th><th>Header 2</th></tr>
            <tr><td>Data 1</td><td>Data 2</td></tr>
        </table>
    </body>
    </html>
    """
    file = tmp_path / "sample.html"
    file.write_text(content, encoding='utf-8')
    return str(file)

@pytest.fixture 
def doc_html_instance():
    # Test data file paths
    html_file_path = "data/test_data/HTML5 Test Page.html"
    # Create an instance of Reader_HTML
    return Doc(path=html_file_path,
          include_images=True,
          actual_first_page=1)

#########DOCX
@pytest.fixture(scope="session")
def sample_docx_path(tmp_path_factory):
    file_path = tmp_path_factory.mktemp("data") / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a test paragraph.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    doc.save(file_path)
    return str(file_path)

@pytest.fixture
def word_reader(sample_docx_path):
    return WordReader(sample_docx_path)
    
@pytest.fixture
def word_reader_with_actual_file():
    docx_file_path = "data/test_data/sample_doc.docx"
    return WordReader(docx_file_path)

@pytest.fixture 
def doc_word_instance():
    # Test data file paths
    docx_file_path = "data/test_data/sample_doc.docx"
    # Create an instance of Reader_HTML
    return Doc(path=docx_file_path,
          include_images=True,
          actual_first_page=1)

#########################XLSX
@pytest.fixture 
def doc_excel_instance():
    # Test data file paths
    excel_file_path = "data/test_data/SampleData.xlsx"
    # Create an instance of Reader_HTML
    return Doc(path=excel_file_path,
          include_images=True,
          actual_first_page=1)

##############PDF
@pytest.fixture 
def doc_pdf_instance():
    # Test data file paths
    pdf_file_path = "data/test_data/pdf-test.pdf"
    # Create an instance of Reader_HTML
    return Doc(path=pdf_file_path,
          include_images=True,
          actual_first_page=1)
