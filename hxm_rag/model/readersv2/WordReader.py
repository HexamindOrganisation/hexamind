from __future__ import (absolute_import, division, print_function,
                        unicode_literals)

import os

import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

class WordReader:
    def __init__(self, path):
        self.path = path
        self.document_structure = self.get_document_structurev3()

    def iter_block_items(self, parent):
        if isinstance(parent, Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Unsupported parent type")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
    
    def table_to_string(self, table):
        table_content = ''
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            table_content += row_text + '\n'
        return table_content.strip()
    
    def get_document_structure(self):
        
        doc = docx.Document(self.path)
        root_container = {'type' : 'container', 'children' : [], 'content' : None, 'level' : 0}
        current_container = root_container


        for block in self.iter_block_items(doc):
            if isinstance(block, Paragraph):
                style = block.style.name if block.style else 'Normal'
                paragraph = block.text.split('\n\n')

                for paragraph_text in paragraph:
                    text = paragraph_text.strip()
                    if not text:
                     continue

                    if style.startswith('Heading'):
                        level = int(style.replace('Heading', ''))
                        while current_container['level'] >= level :
                            current_container = current_container.get('parent', root_container) or root_container
                        
                        new_container = {'type' : 'container', 'children' : [], 'content' : None, 'level' : level, 'parent' : current_container}
                        current_container['children'].append(new_container)
                        current_container = new_container
                    leaf_container = {'type' : 'container', 'children' : [], 'content' : None, 'level' : current_container['level'] + 1, 'parent':current_container}
                    block_content = {'type' : 'block', 'content' : text, 'level' : leaf_container['level'], 'parent' : current_container}
                    leaf_container['children'].append(block_content)
                    current_container['children'].append(leaf_container)
            
            elif isinstance(block, Table):
                table_content = self.table_to_string(block)
                leaf_container = {'type' : 'container', 'children' : [], 'content' : None, 'level' : current_container['level'] + 1, 'parent':current_container}
                table_block = {'type' : 'block', 'content' : table_content, 'level' : leaf_container['level'], 'parent' : current_container}
                leaf_container['children'].append(table_block)
                current_container['children'].append(leaf_container)
        
        def remove_parent_ref(container):
            container.pop('parent', None)
            for child in container.get('children', []):
                if child['type'] == 'container':
                    remove_parent_ref(child)
        
        remove_parent_ref(root_container)
        return root_container
    
    def get_document_structurev2(self):
        doc = docx.Document(self.path)
        root_container = {'type' : 'container', 'children' : [], 'content' : None, 'level' : 0}
        current_container = root_container
        accumulated_text = ''
        last_header_level = 0

        def flush_accumulated_text():
            nonlocal accumulated_text, current_container, last_header_level
            if accumulated_text:
                target_level = last_header_level + (0 if last_header_level else 1)
                parent_container = root_container
                for _ in range(target_level-1):
                    if parent_container['children']:
                        parent_container = parent_container['children'][-1]
                    else:
                        break
            block_content = {'type' : 'block', 'content' : accumulated_text, 'level' : target_level, 'parent' : parent_container}

            if not last_header_level:
                leaf_container = {'type' : 'container', 'children' : [block_content], 'content' : None, 'level' : target_level, 'parent' : parent_container}
                parent_container['children'].append(leaf_container)
            else:
                leaf_container = {'type' : 'container', 'children' : [], 'content' : None, 'level' : target_level, 'parent' : parent_container}
                leaf_container['children'].append(block_content)
                parent_container['children'].append(leaf_container)
            accumulated_text = ''
            last_was_header = False

        for block in self.iter_block_items(doc):
            if isinstance(block, Paragraph):
                style = block.style.name if block.style else 'Normal'
                text = block.text.strip()
            
                if not text:
                    continue

                if style.startswith('Heading'):
                    flush_accumulated_text()
                    level = int(style.replace('Heading', ''))
                    last_header_level = level
                    last_was_header = True
                    accumulated_text = text

                elif text:
                    if accumulated_text:
                        accumulated_text += '\n\n' + text
                    else:
                        accumulated_text = text

            elif isinstance(block, Table):
                table_content = self.table_to_string(block)

                if accumulated_text:
                    accumulated_text += '\n\n' + table_content
                else:
                    accumulated_text = table_content
        
        flush_accumulated_text()

        def remove_parent_ref(container):
            container.pop('parent', None)
            for child in container.get('children', []):
                if child['type'] == 'container':
                    remove_parent_ref(child)
        remove_parent_ref(root_container)
        return root_container


    def get_document_structurev3(self):
        doc = docx.Document(self.path)
        root_container = {'type' : 'container', 'children' : [], 'content' : None, 'level' : 0}
        hierarchy = [root_container]
        accumulated_text = ''
        last_header_level = 0

        def flush_accumulated_text():
            nonlocal accumulated_text
            if accumulated_text:
                block_content = {'type' : 'block', 'content' : accumulated_text, 'level' : last_header_level + 1, 'parent' : hierarchy[-1]}
                leaf_container = {'type' : 'container', 'children' : [block_content], 'content' : None, 'level' : last_header_level + 1, 'parent' : hierarchy[-1]}
                hierarchy[-1]['children'].append(leaf_container)
                accumulated_text = ''
        
        for block in self.iter_block_items(doc):
            if isinstance(block, Paragraph):
                style = block.style.name if block.style else 'Normal'
                text = block.text.strip()

                if not text:
                    continue

                if style.startswith('Heading'):
                    level = int(style.replace('Heading', ''))
                    flush_accumulated_text()

                    while len(hierarchy) > level:
                        hierarchy.pop()
                    last_header_level = level

                    if accumulated_text:
                        accumulated_text += '\n\n' + text
                    else:
                        accumulated_text = text

                else:
                    if accumulated_text:
                        accumulated_text += '\n\n' + text
                    else:
                        accumulated_text = text
            elif isinstance(block, Table):
                table_content = self.table_to_string(block)
                if accumulated_text:
                    accumulated_text += '\n\n' + table_content
                else:
                    accumulated_text = table_content
        flush_accumulated_text()

        def remove_parent_ref(container):
            container.pop('parent', None)
            for child in container.get('children', []):
                if child['type'] == 'container':
                    remove_parent_ref(child)
        remove_parent_ref(root_container)
        return root_container
